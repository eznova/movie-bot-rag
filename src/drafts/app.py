import telebot
import logging
import os
import requests
import chardet
import docx
from langchain_community.llms import HuggingFaceHub
from langchain_community.vectorstores import Chroma
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from langchain.chains import RetrievalQA

# Настройка логирования
logging.basicConfig(format='%(asctime)s - %(name)s - %(levelname)s - %(message)s', level=logging.INFO)
logger = logging.getLogger(__name__)

# Загрузка токенов из переменных окружения
HUGGINGFACE_TOKEN = os.getenv("HUGGINGFACEHUB_API_TOKEN") or "ваш_токен"
BOT_TOKEN = os.getenv("BOT_TOKEN")
TELEGRAM_API_URL = f"https://api.telegram.org/bot{BOT_TOKEN}/"

# Инициализация Telegram-бота
try:
    movie_bot = telebot.TeleBot(BOT_TOKEN)
    logger.info("Бот успешно инициализирован")
except Exception as bot_init_error:
    logger.error("Ошибка инициализации бота: %s", bot_init_error)
    raise

# Инициализация модели и эмбеддингов
llm = HuggingFaceHub(
    repo_id="google/flan-t5-small", 
    model_kwargs={"temperature": 0.7},
    huggingfacehub_api_token=HUGGINGFACE_TOKEN
)
embedder = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")

# Для отслеживания уже приветствованных пользователей
welcomed_users = set()

# Функция для обработки текстовых файлов
def new_gettext(file):
    try:
        if file.endswith('.txt'):
            with open(file, 'rb') as text:
                text_body = text.read()
                enc = chardet.detect(text_body).get("encoding")
                if enc and enc.lower() != "utf-8" and enc.lower() != "windows-1251":
                    text_body = text_body.decode(enc).encode("utf-8").decode("utf-8")
                elif enc and enc.lower() == "windows-1251":
                    with open(file, 'r', encoding='windows-1251') as text:
                        text_body = text.read()
                else:
                    with open(file, 'r', encoding='utf-8') as text:
                        text_body = text.read()
                return text_body
        elif file.endswith('.docx'):
            doc = docx.Document(file)
            text_body = "\n".join(paragraph.text for paragraph in doc.paragraphs)
            return text_body
        else:
            raise ValueError("Unsupported file format")
    except Exception as e:
        logger.error(f"Error processing file: {e}")
        return None

# Функция для разбиения текста на части
def text_chunks(text):
    splitter = RecursiveCharacterTextSplitter(chunk_size=1500, chunk_overlap=100)
    return [Document(page_content=chunk) for chunk in splitter.split_text(text)]

# Функция для отправки сообщений в Telegram
def send_message(chat_id, text):
    url = f"{TELEGRAM_API_URL}sendMessage"
    payload = {
        "chat_id": chat_id,
        "text": text
    }
    response = requests.post(url, data=payload)
    return response.json()

# Функция для отправки просьбы о загрузке файла
def ask_for_file(chat_id):
    send_message(chat_id, "Привет! Пожалуйста, загрузите .txt файл с информацией о фильмах.")

# Функция для получения обновлений (новых сообщений) от Telegram
def get_updates():
    url = f"{TELEGRAM_API_URL}getUpdates"
    response = requests.get(url)
    return response.json()

# Функция для получения файла по file_id
def get_file(file_id):
    url = f"{TELEGRAM_API_URL}getFile"
    payload = {
        "file_id": file_id
    }
    response = requests.get(url, params=payload)
    logger.info(f"Получен файл по file_id: {file_id}, response: {response.json()}")
    return response.json()

# Функция для скачивания файла
def download_file(file_path, chat_id):
    file_url = f"https://api.telegram.org/file/bot{BOT_TOKEN}/{file_path}"
    logger.info(f"Скачивание файла с URL: {file_url}")
    response = requests.get(file_url)
    file_name = f"file_{chat_id}.txt"  # Пример имени файла, можно изменить по желанию
    if response.status_code == 200:
        with open(file_name, 'wb') as f:
            f.write(response.content)
        logger.info(f"Файл {file_name} успешно скачан.")
        
        # Обрабатываем файл
        text = new_gettext(file_name)
        if not text:
            send_message(chat_id, "Не удалось загрузить текст из файла.")
            return
        
        docs = text_chunks(text)

        # Создание базы данных Chroma
        vectorstore = Chroma.from_documents(documents=docs, embedding=embedder)

        # Настройка ретривера
        retriever = vectorstore.as_retriever(search_type="similarity", search_kwargs={"k": 5})

        # Создание цепочки RetrievalQA
        qa_chain = RetrievalQA.from_chain_type(llm=llm, retriever=retriever)

        # Запрос пользователя
        answer = qa_chain.invoke({"query": "Информация о фильмах"})  # Используем invoke вместо run, согласно предупреждению

        # Отправка ответа пользователю
        send_message(chat_id, f"Ответ: {answer}")
    else:
        logger.error(f"Ошибка при скачивании файла. Код ошибки: {response.status_code}")

# Обработчик текстовых сообщений от пользователя
@movie_bot.message_handler(content_types=['text'])
def handle_user_query(message):
    user_id = message['from']['id']
    
    # Если пользователь уже получал приветственное сообщение, пропускаем
    if user_id in welcomed_users:
        return
    
    user_description_query = message.get('text')

    if "/start" in user_description_query:
        welcome_message = "Привет! 👋\n"
        welcome_message += "Я бот, который поможет вам найти фильмы по описаниям! 🎥\n"
        welcome_message += "Загрузите .txt файл с описаниями фильмов, и я помогу вам найти информацию по запросу! 🍿"
        movie_bot.send_message(user_id, welcome_message)
        
        # Добавляем пользователя в список приветствованных
        welcomed_users.add(user_id)
        return

    # Ждем загрузку файла от пользователя
    if 'document' in message:
        file_id = message['document']['file_id']
        file = get_file(file_id)
        if file.get('ok'):
            file_path = file['result']['file_path']
            download_file(file_path, user_id)  # Передаем chat_id в функцию download_file
        else:
            send_message(user_id, "Не удалось получить файл.")
        return

    # Если это не команда /start и не файл, продолжаем поиск по описаниям
    logger.info(f"Получен запрос от пользователя: {user_description_query}")
    send_message(user_id, "Я ожидаю загрузку .txt файла с описаниями фильмов.")

# Главная логика
def main():
    logger.info("Ожидаем сообщений...")
    last_update_id = None

    while True:
        updates = get_updates()
        if updates.get("result"):
            for update in updates["result"]:
                update_id = update["update_id"]
                if update_id != last_update_id:
                    message = update.get("message")
                    if message:  # Проверяем, что сообщение существует
                        handle_user_query(message)
                    last_update_id = update_id

if __name__ == "__main__":
    main()
